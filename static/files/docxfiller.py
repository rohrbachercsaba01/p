from docx import Document
from datetime import datetime, time, timedelta

def calculate_time_difference(start_hour, end_hour):
    start_time = datetime.strptime(start_hour, "%H:%M:%S")
    end_time = datetime.strptime(end_hour, "%H:%M:%S")
    time_difference = end_time - start_time
    if time_difference >= timedelta(hours=6, minutes=30):
        return time_difference - timedelta(minutes=30)
    else: 
        return time_difference

def replace_placeholder(document, placeholder, value):
    for paragraph in document.paragraphs:
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(placeholder, value)

def timedelta_to_string(delta):
    days = delta.days
    hours, remainder = divmod(delta.seconds, 3600)
    minutes, seconds = divmod(remainder, 60)
    
    total_hours = days * 24 + hours
    return "{:02d}:{:02d}:{:02d}".format(total_hours, minutes, seconds)

def fillAttendance(data):

    doc = Document('static/files/doctest.docx')

    tables = doc.tables
    if len(tables) == 0:
        raise ValueError("No tables found in the document.")

    table = tables[0] 
    dt = datetime.strptime("1900-01-01 00:00:00", "%Y-%m-%d %H:%M:%S")
    allHour = dt - datetime.combine(dt.date(), datetime.min.time())

    for person, appointments in data.items():
        for i, appointment in enumerate(appointments):
            row_index = i + 1  
            if row_index < len(table.rows):
                row_cells = table.rows[row_index].cells
                row_cells[0].text = appointment['date']
                row_cells[1].text = appointment['startHour']
                row_cells[2].text = appointment['endHour']
                time_diff = calculate_time_difference(appointment['startHour'], appointment['endHour'])
                row_cells[3].text = str(time_diff)
                allHour = allHour + time_diff
            else:
                break
            lastRow = table.rows[-1].cells
            lastRow[2].text = str(timedelta_to_string(allHour))

        replace_placeholder(doc, "xxxxx", person)
        doc.save('filled_template.docx')
        return 'filled_template.docx'
            